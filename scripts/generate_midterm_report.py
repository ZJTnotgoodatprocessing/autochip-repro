# -*- coding: utf-8 -*-
"""Generate midterm thesis report as Word .docx from JSON content.

Usage:
    python scripts/generate_midterm_report.py
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT = Path(__file__).resolve().parent.parent
CONTENT_JSON = PROJECT / "scripts" / "report_content.json"
DATA_JSON = PROJECT / "outputs" / "verilogeval_both_20260412_173450.json"
BAR_CHART = PROJECT / "outputs" / "reports" / "haiku_pass_rate_bar.png"
RANK_CHART = PROJECT / "outputs" / "reports" / "haiku_per_problem_rank.png"
OUTPUT = PROJECT / "notes" / "midterm.docx"

FONT_BODY = "SimSun"
FONT_HEADING = "SimHei"

# PLACEHOLDER_HELPERS


def set_run_font(run, font_name=FONT_BODY, size=Pt(12), bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.font.bold = bold


def set_cell(cell, text, font_name=FONT_BODY, size=Pt(10), bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    if indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, FONT_BODY, Pt(12))
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, FONT_HEADING)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, FONT_HEADING, Pt(10), True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell(table.rows[r_idx + 1].cells[c_idx], str(val))
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, FONT_BODY, Pt(10))


def add_image(doc, path, width=Inches(4.5)):
    if Path(path).exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# PLACEHOLDER_MAIN

DIFFICULTY = {
    "Prob001_zero": ("Easy", "Comb"),
    "Prob007_wire": ("Easy", "Comb"),
    "Prob014_andgate": ("Easy", "Comb"),
    "Prob024_hadd": ("Easy", "Comb"),
    "Prob027_fadd": ("Easy", "Comb"),
    "Prob031_dff": ("Easy", "Seq"),
    "Prob035_count1to10": ("Easy", "Seq"),
    "Prob041_dff8r": ("Easy", "Seq"),
    "Prob025_reduction": ("Medium", "Comb"),
    "Prob022_mux2to1": ("Medium", "Comb"),
    "Prob050_kmap1": ("Medium", "Comb"),
    "Prob054_edgedetect": ("Medium", "Seq"),
    "Prob068_countbcd": ("Medium", "Seq"),
    "Prob082_lfsr32": ("Medium", "Seq"),
    "Prob085_shift4": ("Medium", "Seq"),
    "Prob030_popcount255": ("Med-Hard", "Comb"),
    "Prob109_fsm1": ("Hard", "FSM"),
    "Prob127_lemmings1": ("Hard", "FSM"),
    "Prob140_fsm_hdlc": ("Hard", "FSM"),
    "Prob144_conwaylife": ("V.Hard", "Seq"),
}


def main():
    C = json.loads(CONTENT_JSON.read_text(encoding="utf-8"))
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    results = data["results"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(C["title"])
    set_run_font(run, FONT_HEADING, Pt(22), True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(C["subtitle"])
    set_run_font(run, FONT_BODY, Pt(14))

    doc.add_paragraph()

    H = C["headings"]
    L = C["table_labels"]

    # === Section 1 ===
    add_heading(doc, H["HEADING_S1"], level=1)
    add_heading(doc, H["HEADING_S1_1"], level=2)
    for t in C["s1_1"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S1_2"], level=2)
    for t in C["s1_2"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S1_3"], level=2)
    for t in C["s1_3"]:
        add_para(doc, t)

    # === Section 2 ===
    add_heading(doc, H["HEADING_S2"], level=1)
    add_heading(doc, H["HEADING_S2_1"], level=2)
    for t in C["s2_1"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S2_2"], level=2)
    for t in C["s2_2"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S2_3"], level=2)
    for t in C["s2_3"]:
        add_para(doc, t)

    # === Section 3 ===
    add_heading(doc, H["HEADING_S3"], level=1)
    add_heading(doc, H["HEADING_S3_1"], level=2)
    for t in C["s3_1"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S3_2"], level=2)
    for t in C["s3_2"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S3_3"], level=2)
    for t in C["s3_3"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S3_4"], level=2)
    for t in C["s3_4"]:
        add_para(doc, t)

    add_heading(doc, H["HEADING_S3_5"], level=2)
    for t in C["s3_5"]:
        add_para(doc, t)

    # === Section 4 ===
    add_heading(doc, H["HEADING_S4"], level=1)
    add_heading(doc, H["HEADING_S4_1"], level=2)
    add_para(doc, C["s4_1"][0])

    # Table 1: Aggregate
    add_table(doc,
        headers=[L["method"], L["pass_count"], L["total"], L["pass_rate"]],
        rows=[
            [L["zs"], "16", "20", "80%"],
            [L["fb"], "18", "20", "90%"],
        ],
    )
    add_caption(doc, L["cap_t1"])

    add_para(doc, C["s4_1"][1])
    add_para(doc, C["s4_1"][2])

    # Figure 1
    add_image(doc, BAR_CHART, Inches(4.5))
    add_caption(doc, L["cap_fig1"])

    add_heading(doc, H["HEADING_S4_2"], level=2)
    add_para(doc, C["s4_2"][0])

    # Table 2: Per-problem detail
    detail_rows = []
    for i, r in enumerate(results, 1):
        name = r["task_name"]
        diff, ctype = DIFFICULTY.get(name, ("?", "?"))
        zs = "PASS" if r["zs_passed"] else "FAIL({:.1%})".format(r["zs_rank"])
        fb = "PASS" if r["fb_passed"] else "FAIL({:.1%})".format(r["fb_rank"])
        iters = str(r["fb_iterations"])
        if r.get("improved"):
            imp = L["yes"]
        elif r["zs_passed"]:
            imp = "-"
        else:
            imp = L["no"]
        detail_rows.append([str(i), name, diff, ctype, zs, fb, iters, imp])

    add_table(doc,
        headers=[L["num"], L["problem"], L["difficulty"], L["type"],
                 L["zs_col"], L["fb_col"], L["iters"], L["improved"]],
        rows=detail_rows,
    )
    add_caption(doc, L["cap_t2"])

    add_para(doc, C["s4_2"][1])
    add_para(doc, C["s4_2"][2])

    # Figure 2
    add_image(doc, RANK_CHART, Inches(6.0))
    add_caption(doc, L["cap_fig2"])

    # 4.3 Case analysis
    add_heading(doc, H["HEADING_S4_3"], level=2)
    for t in C["s4_3"]:
        add_para(doc, t)

    # 4.4 Comparison with original paper
    add_heading(doc, H["HEADING_S4_4"], level=2)
    add_para(doc, C["s4_4"][0])

    # Table 3: Comparison
    add_table(doc,
        headers=[L["col_empty"], L["col_orig"], L["col_ours"]],
        rows=[
            [L["zs_rate"], "~40%", "80%"],
            [L["fb_rate"], "~70%", "90%"],
            [L["abs_gain"], "~30pp", "10pp"],
            [L["max_iter"], "10", "5"],
            [L["k_per_iter"], "1", "3"],
        ],
    )
    add_caption(doc, L["cap_t3"])

    add_para(doc, C["s4_4"][1])

    # === Section 5 ===
    add_heading(doc, H["HEADING_S5"], level=1)
    for t in C["s5"]:
        add_para(doc, t)

    # === Section 6 ===
    add_heading(doc, H["HEADING_S6"], level=1)
    for t in C["s6"]:
        add_para(doc, t)

    # === Section 7 ===
    add_heading(doc, H["HEADING_S7"], level=1)
    for t in C["s7"]:
        add_para(doc, t)

    # === Section 8 ===
    add_heading(doc, H["HEADING_S8"], level=1)
    for t in C["s8"]:
        add_para(doc, t)

    # Save
    doc.save(str(OUTPUT))
    print("Report saved to: {}".format(OUTPUT))

    total_chars = sum(len(p.text) for p in doc.paragraphs)
    print("Approximate character count: {}".format(total_chars))


if __name__ == "__main__":
    main()
