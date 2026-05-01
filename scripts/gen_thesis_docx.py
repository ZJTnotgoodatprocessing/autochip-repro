# -*- coding: utf-8 -*-
"""Generate thesis Word document from Markdown drafts."""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"c:\Users\17819\autochip-repro")
OUTPUT = ROOT / "report" / "thesis" / "thesis_draft_v1.docx"
THESIS_DIR = ROOT / "report" / "thesis"
FIGURES_DIR = ROOT / "report" / "figures"
REPORTS_DIR = ROOT / "outputs" / "reports"

DEPRECATED_FIGS = {"fig_multiturn_comparison.png", "fig_multiturn_matrix.png"}

def read_md(name):
    p = THESIS_DIR / name
    if p.exists():
        return p.read_text(encoding="utf-8")
    return ""

def md_to_paragraphs(md_text):
    result = []
    for line in md_text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("####"):
            result.append((4, s.lstrip("#").strip()))
        elif s.startswith("###"):
            result.append((3, s.lstrip("#").strip()))
        elif s.startswith("##"):
            result.append((2, s.lstrip("#").strip()))
        elif s.startswith("#"):
            result.append((1, s.lstrip("#").strip()))
        elif s.startswith(">") or s.startswith("---"):
            continue
        elif s.startswith("```"):
            continue
        else:
            text = s
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            result.append((0, text))
    return result

def find_figure(text):
    match = re.search(r'(fig_\w+\.png|haiku_\w+\.png)', text)
    if match:
        fname = match.group(1)
        if fname in DEPRECATED_FIGS:
            return None
        for d in [REPORTS_DIR, FIGURES_DIR]:
            p = d / fname
            if p.exists():
                return str(p)
    return None

def add_md_content(doc, md_text):
    figs_in, figs_ph = [], []
    for level, text in md_to_paragraphs(md_text):
        if level > 0:
            doc.add_heading(text, level=min(level, 4))
        elif text.startswith("[") and ("fig_" in text or "haiku_" in text):
            fp = find_figure(text)
            if fp:
                try:
                    doc.add_picture(fp, width=Inches(5.0))
                    figs_in.append(text[:60])
                except Exception:
                    doc.add_paragraph(text)
                    figs_ph.append(text[:60])
            else:
                doc.add_paragraph(text)
                figs_ph.append(text[:60])
        else:
            doc.add_paragraph(text)
    return figs_in, figs_ph

def build_thesis():
    print("Building thesis Word document...")
    doc = Document()
    all_in, all_ph = [], []

    # --- Cover ---
    for _ in range(3):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\u5317\u4eac\u822a\u7a7a\u822a\u5929\u5927\u5b66")
    r.font.size = Pt(26)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\u672c\u79d1\u751f\u6bd5\u4e1a\u8bbe\u8ba1\uff08\u8bba\u6587\uff09")
    r.font.size = Pt(22)
    r.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    cover_items = [
        ("\u8bba\u6587\u9898\u76ee", "\u57fa\u4e8e\u5927\u8bed\u8a00\u6a21\u578b\u4e0eEDA\u5de5\u5177\u53cd\u9988\u7684RTL\u4ee3\u7801\u81ea\u52a8\u751f\u6210\u4e0e\u81ea\u4fee\u590d\u7814\u7a76"),
        ("\u5b66\u9662\u540d\u79f0", "\u8ba1\u7b97\u673a\u5b66\u9662"),
        ("\u4e13    \u4e1a", "\u8ba1\u7b97\u673a\u79d1\u5b66\u4e0e\u6280\u672f"),
        ("\u5b66    \u53f7", "22373311"),
        ("\u5b66\u751f\u59d3\u540d", "\u5f20\u91d1\u6d9b"),
        ("\u6307\u5bfc\u6559\u5e08", "\u738b\u9510"),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        r = p.add_run(label + "\uff1a" + value)
        r.font.size = Pt(14)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026\u5e74 5\u6708")
    r.font.size = Pt(14)

    doc.add_page_break()

    # --- Declaration ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\u672c\u4eba\u58f0\u660e")
    r.font.size = Pt(16)
    r.bold = True

    decl = ("\u672c\u4eba\u90d1\u91cd\u58f0\u660e\uff1a\u6240\u5448\u4ea4\u7684\u6bd5\u4e1a\u8bbe\u8ba1\uff08\u8bba\u6587\uff09\uff0c"
            "\u662f\u672c\u4eba\u5728\u5bfc\u5e08\u6307\u5bfc\u4e0b\u72ec\u7acb\u5b8c\u6210\u7684\u3002"
            "\u5c3d\u6211\u6240\u77e5\uff0c\u9664\u6587\u4e2d\u5df2\u7ecf\u6ce8\u660e\u5f15\u7528\u7684\u5185\u5bb9\u5916\uff0c"
            "\u672c\u6bd5\u4e1a\u8bbe\u8ba1\uff08\u8bba\u6587\uff09\u4e0d\u5305\u542b\u4efb\u4f55\u5176\u4ed6\u4e2a\u4eba\u6216"
            "\u96c6\u4f53\u5df2\u7ecf\u53d1\u8868\u6216\u64b0\u5199\u8fc7\u7684\u4f5c\u54c1\u6210\u679c\u3002"
            "\u5bf9\u672c\u8bba\u6587\u7684\u7814\u7a76\u505a\u51fa\u91cd\u8981\u8d21\u732e\u7684\u4e2a\u4eba\u548c\u96c6\u4f53\uff0c"
            "\u5747\u5df2\u5728\u6587\u4e2d\u4ee5\u660e\u786e\u65b9\u5f0f\u6807\u660e\u3002")
    doc.add_paragraph(decl)
    doc.add_paragraph("")
    doc.add_paragraph("\t\t\u4f5c\u8005\uff1a\u5f20\u91d1\u6d9b")
    doc.add_paragraph("\t\t\u7b7e\u5b57\uff1a")
    doc.add_paragraph("\t\t\u65f6\u95f4\uff1a2026\u5e74 5\u6708")
    doc.add_page_break()

    # --- Chinese Abstract ---
    print("Adding Chinese abstract...")
    add_md_content(doc, read_md("abstract_zh.md"))
    doc.add_page_break()

    # --- English Abstract ---
    print("Adding English abstract...")
    add_md_content(doc, read_md("abstract_en.md"))
    doc.add_page_break()

    # --- TOC placeholder ---
    doc.add_heading("\u76ee       \u5f55", level=1)
    doc.add_paragraph("[\u76ee\u5f55\u5f85\u751f\u6210 - \u8bf7\u5728 Word \u4e2d\u4f7f\u7528\u81ea\u52a8\u76ee\u5f55\u529f\u80fd]")
    doc.add_page_break()

    # --- Chapters 1-6 ---
    ch_files = [
        ("chapter1_introduction.md", "Ch1"),
        ("chapter2_related_work.md", "Ch2"),
        ("chapter3_system_design.md", "Ch3"),
        ("chapter4_experiment_design.md", "Ch4"),
        ("chapter5_results_analysis.md", "Ch5"),
        ("chapter6_conclusion.md", "Ch6"),
    ]
    for fname, label in ch_files:
        print(f"Adding {label}...")
        ins, ph = add_md_content(doc, read_md(fname))
        all_in.extend(ins)
        all_ph.extend(ph)
        doc.add_page_break()

    # --- Acknowledgements ---
    print("Adding acknowledgements...")
    add_md_content(doc, read_md("acknowledgements.md"))
    doc.add_page_break()

    # --- References ---
    print("Adding references...")
    add_md_content(doc, read_md("references.md"))
    doc.add_page_break()

    # --- Appendix ---
    print("Adding appendix...")
    add_md_content(doc, read_md("appendix_plan.md"))

    # --- Save ---
    print(f"Saving to {OUTPUT}...")
    doc.save(str(OUTPUT))
    fsize = OUTPUT.stat().st_size
    print(f"Done! File size: {fsize:,} bytes")

    # --- Verify ---
    print("\n=== Verification ===")
    vdoc = Document(str(OUTPUT))
    print(f"Total paragraphs: {len(vdoc.paragraphs)}")
    full = " ".join(p.text for p in vdoc.paragraphs)
    checks = [
        ("Chinese Abstract", "\u5927\u8bed\u8a00\u6a21\u578b" in full),
        ("English Abstract", "Abstract" in full),
        ("Chapter 1", "\u7eea\u8bba" in full),
        ("Chapter 2", "\u76f8\u5173\u6280\u672f" in full),
        ("Chapter 3", "\u7cfb\u7edf\u8bbe\u8ba1" in full),
        ("Chapter 4", "\u5b9e\u9a8c\u8bbe\u8ba1" in full),
        ("Chapter 5", "\u5b9e\u9a8c\u7ed3\u679c" in full),
        ("Chapter 6", "\u603b\u7ed3" in full),
        ("Acknowledgements", "\u738b\u9510" in full),
        ("References", "AutoChip" in full),
        ("Appendix", "\u9644\u5f55" in full),
    ]
    for name, ok in checks:
        print(f"  {'OK' if ok else 'MISSING'} {name}")

    print(f"\nFigures inserted: {len(all_in)}")
    for f in all_in:
        print(f"  + {f}")
    print(f"Figures as placeholder: {len(all_ph)}")
    for f in all_ph[:15]:
        print(f"  ? {f}")
    if len(all_ph) > 15:
        print(f"  ... and {len(all_ph)-15} more")

if __name__ == "__main__":
    build_thesis()
